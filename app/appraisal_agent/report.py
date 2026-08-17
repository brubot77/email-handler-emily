from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from .models import ActiveDeal


def money(value: Any) -> str:
    if value is None or value == "":
        return "Not available"
    try:
        return f"${float(value):,.0f}"
    except (TypeError, ValueError):
        return str(value)


def number(value: Any, decimals: int = 0) -> str:
    if value is None or value == "":
        return "Not available"
    try:
        return f"{float(value):,.{decimals}f}"
    except (TypeError, ValueError):
        return str(value)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _add_key_value_table(doc: Document, rows: list[tuple[str, str]], widths=(2.4, 4.6)) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        cells[0].width = Inches(widths[0])
        cells[1].width = Inches(widths[1])
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_shading(cells[0], "EDEDED")
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
    doc.add_paragraph()


def _add_section_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def _add_bullets(doc: Document, items: list[str]) -> None:
    if not items:
        doc.add_paragraph("None identified.")
        return
    for item in items:
        doc.add_paragraph(str(item), style="List Bullet")


def _add_comp_table(doc: Document, comps: list[dict[str, Any]], rent: bool = False) -> None:
    if not comps:
        doc.add_paragraph("No sufficiently reliable comparables were documented.")
        return

    if rent:
        headers = ["Address", "Rent", "Beds/Baths", "Sq Ft", "Distance", "Relevance", "Source"]
    else:
        headers = ["Address", "Sale", "Date", "Sq Ft", "$/Sq Ft", "Distance", "Relevance", "Source"]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    _set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        hdr.cells[i].text = header
        _set_cell_shading(hdr.cells[i], "D9E2F3")
        for run in hdr.cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)

    for comp in comps:
        cells = table.add_row().cells
        if rent:
            cells[0].text = str(comp.get("address") or "")
            cells[1].text = money(comp.get("asking_rent"))
            beds = number(comp.get("beds"), 0)
            baths = number(comp.get("baths"), 1)
            cells[2].text = f"{beds} / {baths}"
            cells[3].text = number(comp.get("sqft"), 0)
            cells[4].text = number(comp.get("distance_miles"), 2)
            cells[5].text = str(comp.get("relevance") or "")
            cells[6].text = str(comp.get("source_url") or "")
        else:
            cells[0].text = str(comp.get("address") or "")
            cells[1].text = money(comp.get("sale_price"))
            cells[2].text = str(comp.get("sale_date") or "")
            cells[3].text = number(comp.get("sqft"), 0)
            cells[4].text = money(comp.get("price_per_sqft"))
            cells[5].text = number(comp.get("distance_miles"), 2)
            cells[6].text = str(comp.get("relevance") or "")
            cells[7].text = str(comp.get("source_url") or "")
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(7.5)

    doc.add_paragraph()


def _add_sources(doc: Document, sources: list[dict[str, Any]]) -> None:
    if not sources:
        doc.add_paragraph("No source list returned.")
        return
    for source in sources:
        title = str(source.get("title") or "Source")
        url = str(source.get("url") or "")
        supports = str(source.get("supports") or "")
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(title)
        r.bold = True
        if supports:
            p.add_run(f" - {supports}")
        if url:
            p.add_run(f"\n{url}")


def create_report_docx(deal: ActiveDeal, result: dict[str, Any], output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Arial"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BLU COMPREHENSIVE PROPERTY REVIEW")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run(deal.display_address)
    r.bold = True
    r.font.size = Pt(13)
    review_date = result.get("review_date") or ""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Research date: {review_date}").italic = True

    _add_section_heading(doc, "Executive Summary")
    doc.add_paragraph(str(result.get("executive_summary") or ""))

    appraisal = result.get("appraisal") or {}
    rent = result.get("rent") or {}
    rent_unit = rent.get("per_unit_monthly") or {}
    rent_total = rent.get("total_monthly") or {}

    _add_key_value_table(doc, [
        ("Bank Appraisal - Low", money(appraisal.get("low"))),
        ("Bank Appraisal - Most Likely", money(appraisal.get("most_likely"))),
        ("Bank Appraisal - High", money(appraisal.get("high"))),
        ("Expected Bank Range", f"{money(appraisal.get('expected_bank_range_low'))} - {money(appraisal.get('expected_bank_range_high'))}"),
        ("Appraisal Confidence", str(appraisal.get("confidence") or "")),
        ("Rent / Unit - Most Likely", money(rent_unit.get("most_likely"))),
        ("Total Monthly Rent - Most Likely", money(rent_total.get("most_likely"))),
        ("Recommended Underwriting Rent", money(rent.get("recommended_underwriting_total"))),
        ("Rent Confidence", str(rent.get("confidence") or "")),
        ("Overall Status", str(result.get("status") or "")),
    ])

    _add_section_heading(doc, "Current BLU Deal Information")
    _add_key_value_table(doc, [
        ("Deal", deal.deal or "Not provided"),
        ("Doors", str(deal.doors) if deal.doors is not None else "Not provided"),
        ("Seller Price", money(deal.seller_price)),
        ("Latest Offer", money(deal.latest_offer)),
        ("Rehab Estimate", money(deal.rehab_est)),
        ("Tracker Appraisal Estimate", money(deal.appraisal_est)),
        ("Offer Date", deal.offer_date or "Not provided"),
        ("Offer Status", deal.offer_status or "Not provided"),
        ("Property Notes", deal.property_notes or "Not provided"),
    ])

    subject = result.get("subject") or {}
    _add_section_heading(doc, "Subject Property Details")
    _add_key_value_table(doc, [
        ("Verified Address", str(subject.get("verified_address") or deal.display_address)),
        ("Property Type", str(subject.get("property_type") or "Not available")),
        ("Doors", str(subject.get("doors")) if subject.get("doors") is not None else "Not available"),
        ("Beds / Baths", f"{number(subject.get('bedrooms'), 0)} / {number(subject.get('bathrooms'), 1)}"),
        ("Finished Sq Ft", number(subject.get("sqft"), 0)),
        ("Year Built", number(subject.get("year_built"), 0)),
        ("Lot Sq Ft", number(subject.get("lot_size_sqft"), 0)),
        ("Basement", str(subject.get("basement") or "Not available")),
        ("Garage", str(subject.get("garage") or "Not available")),
        ("Condition", str(subject.get("condition_notes") or "Not available")),
        ("Subject Data Confidence", str(subject.get("data_confidence") or "")),
    ])

    if subject.get("discrepancies"):
        _add_section_heading(doc, "Subject Data Conflicts / Items to Verify", level=2)
        _add_bullets(doc, subject.get("discrepancies") or [])

    _add_section_heading(doc, "Bank Appraisal Analysis")
    _add_key_value_table(doc, [
        ("Valuation Method", str(appraisal.get("valuation_method") or "")),
        ("Low", money(appraisal.get("low"))),
        ("Most Likely", money(appraisal.get("most_likely"))),
        ("High", money(appraisal.get("high"))),
        ("Confidence", str(appraisal.get("confidence") or "")),
        ("Confidence Rationale", str(appraisal.get("confidence_reason") or "")),
    ])

    _add_section_heading(doc, "Appraisal Methodology", level=2)
    doc.add_paragraph(str(appraisal.get("methodology") or ""))
    _add_section_heading(doc, "Adjustment Considerations", level=2)
    doc.add_paragraph(str(appraisal.get("adjustments_summary") or ""))
    _add_section_heading(doc, "Sale Comparables", level=2)
    _add_comp_table(doc, appraisal.get("sale_comps") or [], rent=False)

    for idx, comp in enumerate(appraisal.get("sale_comps") or [], start=1):
        notes = str(comp.get("adjustment_notes") or "").strip()
        if notes:
            p = doc.add_paragraph()
            p.add_run(f"Sale Comp {idx} adjustment notes: ").bold = True
            p.add_run(notes)

    _add_section_heading(doc, "Appraisal Reconciliation", level=2)
    doc.add_paragraph(str(appraisal.get("reconciliation") or ""))

    _add_section_heading(doc, "Rent Analysis")
    _add_key_value_table(doc, [
        ("Rent Basis", str(rent.get("basis") or "")),
        ("Per Unit Range", f"{money(rent_unit.get('low'))} - {money(rent_unit.get('high'))}"),
        ("Per Unit Most Likely", money(rent_unit.get("most_likely"))),
        ("Total Monthly Range", f"{money(rent_total.get('low'))} - {money(rent_total.get('high'))}"),
        ("Total Most Likely", money(rent_total.get("most_likely"))),
        ("Recommended Underwriting Total", money(rent.get("recommended_underwriting_total"))),
        ("Confidence", str(rent.get("confidence") or "")),
        ("Confidence Rationale", str(rent.get("confidence_reason") or "")),
    ])

    _add_section_heading(doc, "Rent Methodology", level=2)
    doc.add_paragraph(str(rent.get("methodology") or ""))
    _add_section_heading(doc, "Rental Comparables", level=2)
    _add_comp_table(doc, rent.get("rent_comps") or [], rent=True)
    _add_section_heading(doc, "Rent Reconciliation", level=2)
    doc.add_paragraph(str(rent.get("reconciliation") or ""))

    _add_section_heading(doc, "Risks / Items to Verify")
    _add_bullets(doc, result.get("risks") or [])

    if result.get("needs_review_reasons"):
        _add_section_heading(doc, "Needs Review")
        _add_bullets(doc, result.get("needs_review_reasons") or [])

    _add_section_heading(doc, "Research Sources")
    _add_sources(doc, result.get("research_sources") or [])

    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = disclaimer.add_run(
        "This is an internal forecast for acquisition/underwriting use and is not a licensed appraisal, broker price opinion, or guarantee of value or rent."
    )
    rr.italic = True
    rr.font.size = Pt(7.5)

    doc.save(output_path)
    return output_path
